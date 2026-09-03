import io
import unittest
import docx
import pypdf
from app import create_app
from extensions import db
from models.user import User
from models.organisatie import Organisatie, UserOrganisatie
from models.document import Folder, Document
from utils.text_extractor import extraheer_tekst_uit_bestand, genereer_zoek_snippet


class TestDocumentSearch(unittest.TestCase):
    def setUp(self):
        self.app = create_app()
        self.app.config['TESTING'] = True
        self.app.config['WTF_CSRF_ENABLED'] = False
        self.client = self.app.test_client()

        with self.app.app_context():
            # Gebruik bestaande of maak testorganisatie en gebruiker
            org = Organisatie.query.filter_by(id=1).first()
            if not org:
                org = Organisatie(id=1, naam='Test Organisatie', slug='test-org', actief=True)
                db.session.add(org)
                db.session.commit()
            self.org_id = org.id

            user = User.query.filter_by(naam='TestUserDoc').first()
            if not user:
                user = User(naam='TestUserDoc', email='testdoc@example.com', wachtwoord_hash='dummy', rol='beheerder')
                db.session.add(user)
                db.session.commit()
                uo = UserOrganisatie(user_id=user.id, organisatie_id=org.id, rol='beheerder', actief=True)
                db.session.add(uo)
                db.session.commit()
            self.user_id = user.id

    def tearDown(self):
        with self.app.app_context():
            # Ruim gemaakte testdocumenten op
            Document.query.filter(Document.bestandsnaam.like('test_search_%')).delete()
            db.session.commit()

    def test_text_extractor_formats(self):
        """Test tekstextractie uit diverse formaten."""
        # 1. Tekstbestand
        txt_content = "Dit is een geheime inhoud voor de digidokters handleiding.".encode('utf-8')
        t1 = extraheer_tekst_uit_bestand(txt_content, 'bestand.txt', 'text/plain')
        self.assertIn("geheime inhoud", t1)

        # 2. Word (.docx)
        doc = docx.Document()
        doc.add_paragraph("Informatie over itsme en identiteitskaart.")
        buf = io.BytesIO()
        doc.save(buf)
        t2 = extraheer_tekst_uit_bestand(buf.getvalue(), 'handleiding.docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        self.assertIn("itsme en identiteitskaart", t2)

        # 3. PDF
        writer = pypdf.PdfWriter()
        writer.add_blank_page(width=72, height=72)
        buf_pdf = io.BytesIO()
        writer.write(buf_pdf)
        t3 = extraheer_tekst_uit_bestand(buf_pdf.getvalue(), 'doc.pdf', 'application/pdf')
        self.assertIsNone(t3)  # Blanke pagina heeft geen tekst

        # 4. Niet-ondersteund (bijv. png)
        t4 = extraheer_tekst_uit_bestand(b'\x89PNG\r\n\x1a\n', 'foto.png', 'image/png')
        self.assertIsNone(t4)

    def test_snippet_generator(self):
        """Test genereren van zoeksnippets."""
        tekst = "Hier is het begin van de tekst. De burger had een probleem met itsme activatie op zijn telefoon. Dit is het einde."
        snippet = genereer_zoek_snippet(tekst, "itsme", context_lengte=20)
        self.assertIsNotNone(snippet)
        self.assertIn("itsme", snippet)
        self.assertTrue(snippet.startswith("..."))
        self.assertTrue(snippet.endswith("..."))

        # Niet bestaand zoekwoord
        self.assertIsNone(genereer_zoek_snippet(tekst, "onbestaandewoordxyz"))

    def test_upload_and_content_search(self):
        """Test dat een geüpload bestand via de zoekbalk gevonden wordt op basis van inhoud."""
        with self.client.session_transaction() as sess:
            sess['_user_id'] = str(self.user_id)
            sess['organisatie_id'] = self.org_id

        # Maak een Word document met een unieke term in de inhoud
        unique_term = "XyZzY999SpecialKey"
        doc = docx.Document()
        doc.add_paragraph(f"Informatie over {unique_term} activatie voor vrijwilligers.")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        # Upload document
        upload_resp = self.client.post('/documenten/upload', data={
            'bestand': (buf, 'test_search_document.docx'),
            'omschrijving': 'Algemene omschrijving zonder de unieke term'
        }, follow_redirects=True)
        self.assertEqual(upload_resp.status_code, 200)

        # Controleer dat document in de database tekst_inhoud heeft
        with self.app.app_context():
            doc_record = Document.query.filter_by(bestandsnaam='test_search_document.docx').first()
            self.assertIsNotNone(doc_record)
            self.assertIsNotNone(doc_record.tekst_inhoud)
            self.assertIn(unique_term, doc_record.tekst_inhoud)

        # Zoek via zoekbalk op de unieke term (die NIET in naam of omschrijving staat)
        search_resp = self.client.get(f'/documenten/?zoek={unique_term}')
        self.assertEqual(search_resp.status_code, 200)
        html = search_resp.data.decode('utf-8')

        # Document moet in zoekresultaten verschijnen
        self.assertIn('test_search_document.docx', html)
        self.assertIn('Gevonden in inhoud', html)
        self.assertIn(unique_term, html)


if __name__ == '__main__':
    unittest.main()
