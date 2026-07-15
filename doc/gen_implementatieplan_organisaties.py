"""Genereer Word-document: implementatieplan multi-organisatie."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


def _heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def _para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def _bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    if level:
        p.paragraph_format.left_indent = Cm(1.27 * level)
    return p


def _numbered(doc, text):
    return doc.add_paragraph(text, style='List Number')


def _table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()
    return table


def build_document():
    doc = Document()

    # Titelpagina
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Digidokters\n')
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    sub = title.add_run('Implementatieplan: multi-organisatie (multi-tenant)')
    sub.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('\nVersie 1.0 — juli 2026\n').italic = True
    meta.add_run('Status: denkoefening / ontwerpvoorstel\n').italic = True
    meta.add_run('Nog niet geïmplementeerd').italic = True

    doc.add_page_break()

    # 1. Inleiding
    _heading(doc, '1. Inleiding')
    _para(doc,
        'Digidokters is momenteel een single-tenant applicatie: één organisatie, één '
        'gebruikersbestand, één set registraties. Andere organisaties tonen interesse om '
        'dezelfde applicatie te gebruiken. Dit document beschrijft een implementatieplan '
        'om de app uit te breiden met het concept organisatie, waarbij alle data in '
        'dezelfde database en dezelfde tabellen blijft (shared database, shared schema).')
    _para(doc,
        'Kernprincipe: voor de eindgebruiker moeten de wijzigingen grotendeels transparant '
        'zijn. Medewerkers zien enkel data van hun eigen organisatie; nieuwe registraties '
        'worden automatisch aan die organisatie gekoppeld.')

    _heading(doc, '1.1 Doelstellingen', level=2)
    for item in [
        'Meerdere organisaties ondersteunen in één applicatie-instantie.',
        'Strikte datascheiding per organisatie (geen lekken tussen tenants).',
        'Gebruikers kunnen aan één of meerdere organisaties gekoppeld worden.',
        'Bestaande organisatie (huidige productiedata) migreren zonder verlies.',
        'Minimale impact op de bestaande UI voor gebruikers met één organisatie.',
    ]:
        _bullet(doc, item)

    _heading(doc, '1.2 Niet-doelstellingen (voorlopig)', level=2)
    for item in [
        'Aparte database per organisatie.',
        'Aparte subdomain/URL per organisatie (optioneel later).',
        'White-label branding per organisatie.',
        'Facturatie of licentiebeheer.',
    ]:
        _bullet(doc, item)

    # 2. Architectuurkeuze
    _heading(doc, '2. Architectuurkeuze')
    _para(doc,
        'Gekozen model: row-level multi-tenancy. Elke tenant (organisatie) wordt '
        'geïdentificeerd via een foreign key organisatie_id op de relevante tabellen. '
        'Dit is de meest gangbare aanpak voor SaaS-toepassingen van deze schaal.')

    _table(doc,
        ['Aspect', 'Keuze', 'Motivatie'],
        [
            ['Databasemodel', 'Shared schema + organisatie_id', 'Eenvoudig te beheren, één deploy, lage kosten'],
            ['Tenant-isolatie', 'Applicatielaag (query filtering)', 'Geen aparte DB-per-tenant nodig'],
            ['Gebruikers', 'Globaal account + koppeltabel', 'Eén login voor meerdere organisaties'],
            ['Rollen', 'Per organisatie-lidmaatschap', 'Beheerder in org A, medewerker in org B'],
            ['Referentiedata', 'Per organisatie', 'Elke org heeft eigen Digidokters, categorieën, toestellen'],
        ])

    _para(doc,
        'Alternatief (niet gekozen): schema-per-tenant of database-per-tenant. Dat biedt '
        'strengere isolatie maar verhoogt operationele complexiteit aanzienlijk en past '
        'minder bij de huidige Render/Supabase-setup.')

    # 3. Datamodel
    _heading(doc, '3. Datamodel')

    _heading(doc, '3.1 Nieuwe tabellen', level=2)

    _heading(doc, 'organisaties', level=3)
    _table(doc,
        ['Kolom', 'Type', 'Toelichting'],
        [
            ['id', 'INTEGER PK', 'Unieke identifier'],
            ['naam', 'VARCHAR(150)', 'Weergavenaam, bv. "Digidokters Leuven"'],
            ['slug', 'VARCHAR(80) UNIQUE', 'URL-vriendelijke code, bv. "leuven"'],
            ['actief', 'BOOLEAN', 'Gedeactiveerde org: geen toegang'],
            ['aangemaakt_op', 'DATETIME', 'Audit'],
        ])

    _heading(doc, 'user_organisaties (koppeltabel)', level=3)
    _para(doc,
        'Many-to-many relatie tussen users en organisaties, met rol per lidmaatschap.')
    _table(doc,
        ['Kolom', 'Type', 'Toelichting'],
        [
            ['id', 'INTEGER PK', ''],
            ['user_id', 'FK → users.id', 'Gebruiker'],
            ['organisatie_id', 'FK → organisaties.id', 'Organisatie'],
            ['rol', 'VARCHAR(20)', "'beheerder' | 'medewerker' — per organisatie"],
            ['actief', 'BOOLEAN', 'Lidmaatschap actief/inactief'],
            ['aangemaakt_op', 'DATETIME', ''],
        ])
    _para(doc, 'Unieke constraint: (user_id, organisatie_id).')

    _heading(doc, '3.2 Bestaande tabellen — nieuwe kolom organisatie_id', level=2)
    _table(doc,
        ['Tabel', 'organisatie_id', 'Extra wijzigingen'],
        [
            ['registrations', 'NOT NULL FK', 'Uniek: (organisatie_id, registratienummer) i.p.v. globaal uniek'],
            ['digidokters', 'NOT NULL FK', 'Naam uniek per organisatie'],
            ['age_categories', 'NOT NULL FK', 'Naam uniek per organisatie'],
            ['devices', 'NOT NULL FK', 'Naam uniek per organisatie'],
            ['users', '—', 'Geen organisatie_id; koppeling via user_organisaties'],
        ])

    _heading(doc, '3.3 Wijziging users-tabel', level=2)
    _para(doc,
        'Gebruikers blijven globaal (één login). De kolom rol op users verdwijnt op termijn '
        'ten gunste van de rol in user_organisaties. Tijdens migratie: rol op users behouden '
        'als fallback voor backwards compatibility, daarna verwijderen.')
    _para(doc,
        'Login blijft op gebruikersnaam (naam). Uniciteit van naam blijft globaal — '
        'twee organisaties kunnen geen gebruiker met dezelfde login hebben. Dit is '
        'acceptabel en voorkomt verwarring bij multi-org gebruikers.')

    _heading(doc, '3.4 Entity-relationship (conceptueel)', level=2)
    _para(doc,
        'organisaties ←── user_organisaties ──→ users\n'
        'organisaties ←── registrations\n'
        'organisaties ←── digidokters, age_categories, devices\n'
        'registrations → digidokter, leeftijdscategorie, toestel (allemaal binnen dezelfde org)')

    _para(doc,
        'Belangrijk: foreign keys naar digidokters/leeftijd/toestel moeten binnen dezelfde '
        'organisatie blijven. Dit wordt afgedwongen in de applicatielaag (validatie bij '
        'opslaan) en optioneel via database-check constraints.')

    # 4. Tenant context
    _heading(doc, '4. Tenant context — hoe datascheiding werkt')

    _heading(doc, '4.1 Huidige organisatie in sessie', level=2)
    _para(doc,
        'Na login wordt de actieve organisatie opgeslagen in de Flask-sessie '
        '(session["organisatie_id"]). Alle queries filteren op deze waarde.')

    _numbered(doc, 'Gebruiker logt in.')
    _numbered(doc, 'Systeem laadt organisaties via user_organisaties.')
    _numbered(doc, 'Heeft gebruiker precies één actieve org → automatisch selecteren (transparant).')
    _numbered(doc, 'Heeft gebruiker meerdere orgs → org kiezen (eenmalig) of org-switcher tonen.')
    _numbered(doc, 'Alle routes lezen/schrijven enkel data waar organisatie_id = huidige org.')

    _heading(doc, '4.2 Centrale query-helper', level=2)
    _para(doc,
        'Introduceer een utility-module utils/tenant.py met functies zoals:')
    for item in [
        'get_huidige_organisatie() — haalt org uit sessie, valideert lidmaatschap',
        'filter_op_organisatie(query, Model) — voegt .filter(Model.organisatie_id == ...) toe',
        'set_organisatie_id_op_model(instance) — zet FK bij aanmaken',
        'require_organisatie — decorator/before_request guard',
    ]:
        _bullet(doc, item)

    _para(doc,
        'Optioneel (fase 2): SQLAlchemy event listener of custom Query class die automatisch '
        'filtert. Voordeel: minder risico om een query te vergeten. Nadeel: complexer debuggen. '
        'Aanbeveling: start met expliciete helper, evalueer later auto-filter.')

    _heading(doc, '4.3 before_request hook', level=2)
    _para(doc,
        'In app.py: voor elke authenticated request controleren of organisatie_id in sessie '
        'geldig is en gebruiker nog lid is. Zo niet → redirect naar org-selectie of logout.')

    # 5. Rollen en rechten
    _heading(doc, '5. Rollen en rechten')

    _table(doc,
        ['Rol', 'Scope', 'Rechten'],
        [
            ['Platformbeheerder (nieuw, optioneel)', 'Hele applicatie', 'Organisaties aanmaken/beheren, globaal overzicht'],
            ['Beheerder', 'Per organisatie', 'Gebruikers org, keuzelijsten, import, alles medewerker'],
            ['Medewerker', 'Per organisatie', 'Registraties CRUD, export, statistieken'],
        ])

    _para(doc,
        'De huidige admin_required decorator wordt uitgebreid tot org_admin_required die '
        'zowel rol "beheerder" in de huidige org als actief lidmaatschap controleert.')

    _para(doc,
        'Gebruikersbeheer wijzigt: een org-beheerder ziet enkel gebruikers die lid zijn van '
        'zijn organisatie. Nieuwe gebruiker aanmaken = globaal account + automatisch '
        'lidmaatschap in huidige org. Bestaande gebruiker toevoegen = lidmaatschap aanmaken '
        'voor user die al bestaat (uitnodigingsflow, fase 2).')

    # 6. Impact per module
    _heading(doc, '6. Impact per module')

    modules = [
        ('models/', 'Nieuw model Organisatie, UserOrganisatie; organisatie_id op 4 tabellen'),
        ('routes/registrations.py', 'Alle queries filteren; registratienummer per org; FK-validatie'),
        ('routes/admin.py', 'Scoped naar org; gebruikersbeheer via koppeltabel; rol per org'),
        ('routes/stats.py', 'Alle aggregaties filteren op organisatie_id'),
        ('routes/import_export.py', 'Import/export scoped; geen cross-org data'),
        ('utils/import_handler.py', 'organisatie_id meegeven; dubbele detectie per org'),
        ('utils/export_handler.py', 'Filter op organisatie_id'),
        ('utils/decorators.py', 'org_admin_required, org_member_required'),
        ('routes/auth.py', 'Na login org-selectie; sessie organisatie_id'),
        ('app.py seed', 'Default org aanmaken; bestaande seed-data koppelen'),
        ('templates/base.html', 'Org-switcher (alleen bij multi-org); org-naam in sidebar'),
        ('migrations/', 'Data-migratie: default org + backfill organisatie_id'),
    ]
    _table(doc, ['Module', 'Wijziging'], modules)

    # 7. UX
    _heading(doc, '7. Gebruikerservaring')

    _heading(doc, '7.1 Eén organisatie (meeste gebruikers)', level=2)
    _para(doc, 'Geen zichtbare wijziging. Na login werkt alles identiek. Geen org-switcher.')

    _heading(doc, '7.2 Meerdere organisaties', level=2)
    _para(doc,
        'Discrete org-switcher in de sidebar (dropdown). Wisselen = sessie update + redirect '
        'naar registratielijst. Duidelijke org-naam in header voorkomt verwarring.')

    _heading(doc, '7.3 Org-selectiescherm', level=2)
    _para(doc,
        'Nieuw scherm /kies-organisatie na login als gebruiker nog geen actieve org in sessie '
        'heeft. Kaarten met org-namen; klik = selecteren en doorgaan.')

    # 8. Migratie bestaande data
    _heading(doc, '8. Migratiestrategie bestaande data')

    _numbered(doc, 'Maak tabel organisaties aan.')
    _numbered(doc, 'Insert default organisatie, bv. naam="Digidokters", slug="digidokters".')
    _numbered(doc, 'Voeg organisatie_id toe aan registrations, digidokters, age_categories, devices (nullable eerst).')
    _numbered(doc, 'Backfill: UPDATE alle rijen SET organisatie_id = <default org id>.')
    _numbered(doc, 'Maak organisatie_id NOT NULL.')
    _numbered(doc, 'Pas unique constraints aan (registratienummer per org).')
    _numbered(doc, 'Maak user_organisaties; migreer bestaande users met hun huidige rol naar default org.')
    _numbered(doc, 'Verwijder kolom users.rol (aparte migratie na stabilisatie).')

    _para(doc,
        'Rollback-plan: migraties in stappen, backup vóór productie-deploy (bestaande '
        'wekelijkse pg_dump GitHub Action).')

    # 9. Implementatiefasen
    _heading(doc, '9. Implementatiefasen')

    _heading(doc, 'Fase 1 — Fundament (1–2 weken)', level=2)
    for item in [
        'Models Organisatie + UserOrganisatie',
        'Database-migratie + backfill default org',
        'utils/tenant.py + sessie-organisatie_id',
        'before_request validatie',
        'Registraties-module volledig scoped',
    ]:
        _bullet(doc, item)
    _para(doc, 'Resultaat: bestaande org werkt ongewijzigd; technische basis staat.')

    _heading(doc, 'Fase 2 — Alle modules scoped (1 week)', level=2)
    for item in [
        'Admin, stats, import/export scoped',
        'Registratienummer-generatie per org',
        'Decorators org_admin_required',
        'Seed-commando aanpassen',
    ]:
        _bullet(doc, item)

    _heading(doc, 'Fase 3 — Multi-org UX (3–5 dagen)', level=2)
    for item in [
        'Org-selectiescherm na login',
        'Org-switcher in sidebar',
        'Gebruikersbeheer via koppeltabel',
    ]:
        _bullet(doc, item)

    _heading(doc, 'Fase 4 — Nieuwe organisaties onboarding (3–5 dagen)', level=2)
    for item in [
        'Platformbeheerder: organisatie aanmaken',
        'Automatisch seed keuzelijsten bij nieuwe org',
        'Eerste beheerder aanmaken per org',
        'Documentatie voor onboarding nieuwe org',
    ]:
        _bullet(doc, item)

    _heading(doc, 'Fase 5 — Hardening & tests (1 week)', level=2)
    for item in [
        'Integratietests: datascheiding (org A ziet org B niet)',
        'Test multi-org gebruiker + switcher',
        'Security review cross-tenant toegang',
        'users.rol kolom verwijderen',
    ]:
        _bullet(doc, item)

    # 10. Risico's
    _heading(doc, '10. Risico\'s en mitigaties')

    risks = [
        ('Vergeten org-filter in query', 'Hoog', 'Centrale helper; code review checklist; tests'),
        ('Registratienummer collision', 'Medium', 'Composite unique (org_id, nummer)'),
        ('FK cross-org (verkeerde digidokter_id)', 'Medium', 'Validatie bij save: gekozen FK hoort bij zelfde org'),
        ('Gebruiker zonder org', 'Medium', 'Blokkeer toegang; duidelijke foutmelding'),
        ('Migratie productiedata', 'Hoog', 'Backup + staging test + nullable-then-backfill patroon'),
        ('Performance bij groei', 'Laag', 'Index op organisatie_id; composite indexes waar nodig'),
    ]
    _table(doc, ['Risico', 'Ernst', 'Mitigatie'], risks)

    # 11. Open beslissingen
    _heading(doc, '11. Open beslissingen (voor implementatie)')

    decisions = [
        ('Platformbeheerder nodig?', 'Ja, als jij orgs beheert; nee als elke org zelfstandig onboard via script'),
        ('Uitnodigingsflow e-mail?', 'Fase 2: bestaande user toevoegen aan org via e-mail'),
        ('Subdomain per org?', 'Nu niet; later mogelijk via slug in URL'),
        ('Gedeelde standaard keuzelijsten?', 'Template bij nieuwe org (kopie seed-data)'),
        ('users.rol verwijderen wanneer?', 'Na fase 5, als alles via user_organisaties loopt'),
        ('Login op e-mail i.p.v. naam?', 'Optioneel later; nu naam behouden'),
    ]
    _table(doc, ['Vraag', 'Opties / aanbeveling'], decisions)

    # 12. Geschatte inspanning
    _heading(doc, '12. Geschatte totale inspanning')
    _table(doc,
        ['Fase', 'Duur', 'Cumulatief'],
        [
            ['Fase 1 — Fundament', '1–2 weken', '1–2 weken'],
            ['Fase 2 — Modules scoped', '1 week', '2–3 weken'],
            ['Fase 3 — Multi-org UX', '3–5 dagen', '3–4 weken'],
            ['Fase 4 — Onboarding', '3–5 dagen', '4–5 weken'],
            ['Fase 5 — Hardening', '1 week', '5–6 weken'],
        ])
    _para(doc,
        'Schatting voor één ontwikkelaar, parttime. Parallel testen op staging aanbevolen.')

    # 13. Conclusie
    _heading(doc, '13. Conclusie')
    _para(doc,
        'Multi-organisatie via organisatie_id en een koppeltabel user_organisaties is een '
        'beproefde, haalbare uitbreiding die past bij de huidige architectuur. De impact is '
        'voornamelijk in de datalaag (elke query filteren) en de authenticatieflow (org-context '
        'in sessie). Voor bestaande gebruikers met één organisatie blijft de ervaring '
        'identiek.')
    _para(doc,
        'Aanbevolen start: Fase 1 op een feature branch, migratie testen tegen een kopie van '
        'de productiedatabase, daarna geleidelijk uitrollen.')

    return doc


def main():
    out = Path(__file__).parent / 'Implementatieplan_multi-organisatie.docx'
    doc = build_document()
    doc.save(out)
    print(f'Geschreven: {out}')


if __name__ == '__main__':
    main()
