"""
Utility voor tekstextractie uit diverse documentformaten en het genereren van zoeksnippets.
Ondersteunt: PDF, Word (.docx), Excel (.xlsx), tekstbestanden (.txt, .csv, .md, enz.).
"""
import io
import os
import re


def extraheer_tekst_uit_bestand(inhoud_bytes, bestandsnaam='', mime_type=''):
    """
    Extraheert platte tekst uit binaire bestandsdata op basis van extensie of mime-type.
    Geeft een schone string terug, of None als extractie niet van toepassing/mislukt is.
    """
    if not inhoud_bytes:
        return None

    ext = os.path.splitext(bestandsnaam)[1].lower().lstrip('.') if bestandsnaam else ''
    mime_type = (mime_type or '').lower()

    # 1. PDF
    if ext == 'pdf' or 'application/pdf' in mime_type:
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(inhoud_bytes))
            pagina_teksten = []
            for pagina in reader.pages:
                t = pagina.extract_text()
                if t:
                    pagina_teksten.append(t)
            if pagina_teksten:
                return _opschonen_tekst('\n'.join(pagina_teksten))
            return None
        except Exception:
            return None

    # 2. Word (.docx)
    if ext in ('docx', 'docm') or 'word' in mime_type:
        try:
            import docx
            doc = docx.Document(io.BytesIO(inhoud_bytes))
            delen = []
            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    delen.append(p.text.strip())
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            delen.append(cell.text.strip())
            if delen:
                return _opschonen_tekst('\n'.join(delen))
            return None
        except Exception:
            return None

    # 3. Excel (.xlsx, .xlsm)
    if ext in ('xlsx', 'xlsm') or 'excel' in mime_type or 'spreadsheet' in mime_type:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(inhoud_bytes), data_only=True, read_only=True)
            delen = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    cel_teksten = [str(c).strip() for c in row if c is not None and str(c).strip()]
                    if cel_teksten:
                        delen.append(' '.join(cel_teksten))
            if delen:
                return _opschonen_tekst('\n'.join(delen))
            return None
        except Exception:
            return None

    # 4. Platte tekstbestanden (.txt, .csv, .md, .tsv, .json, .xml, .html)
    if ext in ('txt', 'csv', 'md', 'tsv', 'json', 'xml', 'html', 'log') or 'text/' in mime_type or 'json' in mime_type:
        try:
            # Probeer UTF-8, daarna latin-1 fallback
            try:
                tekst = inhoud_bytes.decode('utf-8')
            except UnicodeDecodeError:
                tekst = inhoud_bytes.decode('latin-1', errors='ignore')
            return _opschonen_tekst(tekst)
        except Exception:
            return None

    return None


def _opschonen_tekst(tekst):
    """Vervangt overtollige witruimtes en regels door nette spaties/enters."""
    if not tekst:
        return None
    # Verwijder null bytes
    tekst = tekst.replace('\x00', '')
    # Vervang 3+ nieuwe regels door dubbele nieuwe regel
    tekst = re.sub(r'\n{3,}', '\n\n', tekst)
    tekst = tekst.strip()
    return tekst if tekst else None


def genereer_zoek_snippet(tekst, zoekterm, context_lengte=60):
    """
    Genereert een kort contextfragment (snippet) rondom de eerste vindplaats van zoekterm.
    Geeft None terug als de zoekterm niet in de tekst voorkomt.
    """
    if not tekst or not zoekterm:
        return None

    zoekterm_schoon = zoekterm.strip()
    if not zoekterm_schoon:
        return None

    # Case-insensitieve zoekopdracht
    idx = tekst.lower().find(zoekterm_schoon.lower())
    if idx == -1:
        return None

    start = max(0, idx - context_lengte)
    eind = min(len(tekst), idx + len(zoekterm_schoon) + context_lengte)

    snippet = tekst[start:eind].replace('\n', ' ').strip()
    # Verwijder meervoudige spaties
    snippet = re.sub(r'\s+', ' ', snippet)

    prefix = '...' if start > 0 else ''
    suffix = '...' if eind < len(tekst) else ''

    return f"{prefix}{snippet}{suffix}"
